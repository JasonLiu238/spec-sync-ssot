"""
Spec-Sync SSOT - Flask API Server
RESTful API for frontend
"""

from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from flask_socketio import SocketIO, emit
import os
import sys
import yaml
import json
import subprocess
from datetime import datetime
from pathlib import Path
import logging
import re

# Add project root to Python path
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize Flask
app = Flask(__name__)
app.config['SECRET_KEY'] = 'spec-sync-ssot-secret-key-2025'
CORS(app)  # Enable CORS
socketio = SocketIO(app, cors_allowed_origins="*")

# Setup paths
SSOT_DIR = project_root / 'ssot'
MAPPING_DIR = project_root / 'mapping'
TEMPLATES_DIR = project_root / 'templates'
OUTPUT_DIR = project_root / 'output'

# Ensure directories exist
for dir_path in [SSOT_DIR, MAPPING_DIR, TEMPLATES_DIR, OUTPUT_DIR]:
    dir_path.mkdir(exist_ok=True)


# ============================================================================
# Helpers: SSOT access + Token scan/replace
# ============================================================================

def _load_ssot() -> dict:
    ssot_file = SSOT_DIR / 'master.yaml'
    if not ssot_file.exists():
        raise FileNotFoundError('SSOT 檔案不存在')
    with open(ssot_file, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f) or {}


def _get_nested_value(data: dict, path: str):
    cur = data
    try:
        for key in path.split('.'):
            cur = cur[key]
        return cur
    except Exception:
        return None


_TOKEN_REGEX = re.compile(r"\{([A-Za-z0-9_.-]+)\}")


def _scan_tokens_docx(path: Path) -> set[str]:
    from docx import Document  # type: ignore
    tokens: set[str] = set()
    doc = Document(str(path))

    def collect(paragraph):
        text = paragraph.text or ''
        for m in _TOKEN_REGEX.finditer(text):
            tokens.add(m.group(1))

    for p in doc.paragraphs:
        collect(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    collect(p)
    return tokens


def _replace_tokens_docx(path: Path, out_path: Path, ssot: dict) -> dict:
    from docx import Document  # type: ignore
    missing: set[str] = set()
    replaced: dict[str, str] = {}
    doc = Document(str(path))

    def replace(paragraph):
        text = paragraph.text or ''
        def repl(m):
            key = m.group(1)
            val = _get_nested_value(ssot, key)
            if val is None:
                missing.add(key)
                return m.group(0)
            s = str(val)
            replaced[key] = s
            return s
        new_text = _TOKEN_REGEX.sub(repl, text)
        if new_text != text:
            paragraph.text = new_text

    for p in doc.paragraphs:
        replace(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace(p)

    doc.save(str(out_path))
    return {"missing": sorted(missing), "replaced": replaced}


def _scan_tokens_xlsx(path: Path) -> set[str]:
    from openpyxl import load_workbook  # type: ignore
    tokens: set[str] = set()
    wb = load_workbook(str(path), data_only=False)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    for m in _TOKEN_REGEX.finditer(cell.value):
                        tokens.add(m.group(1))
    return tokens


def _replace_tokens_xlsx(path: Path, out_path: Path, ssot: dict) -> dict:
    from openpyxl import load_workbook  # type: ignore
    missing: set[str] = set()
    replaced: dict[str, str] = {}
    wb = load_workbook(str(path), data_only=False)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    original = cell.value
                    def repl(m):
                        key = m.group(1)
                        val = _get_nested_value(ssot, key)
                        if val is None:
                            missing.add(key)
                            return m.group(0)
                        s = str(val)
                        replaced[key] = s
                        return s
                    new_val = _TOKEN_REGEX.sub(repl, original)
                    if new_val != original:
                        cell.value = new_val
    wb.save(str(out_path))
    return {"missing": sorted(missing), "replaced": replaced}


# ============================================================================
# API: SSOT 資料管理
# ============================================================================

@app.route('/api/ssot', methods=['GET'])
def get_ssot():
    """讀取 SSOT 資料"""
    try:
        ssot_file = SSOT_DIR / 'master.yaml'
        if not ssot_file.exists():
            return jsonify({'error': 'SSOT 檔案不存在'}), 404
        
        with open(ssot_file, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
        
        return jsonify({
            'success': True,
            'data': data,
            'last_modified': datetime.fromtimestamp(ssot_file.stat().st_mtime).isoformat()
        })
    except Exception as e:
        logger.error(f"讀取 SSOT 失敗: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/ssot', methods=['POST'])
def update_ssot():
    """更新 SSOT 資料"""
    try:
        data = request.get_json()
        
        # 驗證資料結構
        if not data:
            return jsonify({'error': '無效的資料'}), 400
        
        ssot_file = SSOT_DIR / 'master.yaml'
        
        # 備份現有檔案
        if ssot_file.exists():
            backup_file = SSOT_DIR / f'master.yaml.backup.{int(datetime.now().timestamp())}'
            import shutil
            shutil.copy(ssot_file, backup_file)
        
        # 寫入新資料
        data['last_updated'] = datetime.now().strftime('%Y-%m-%d')
        with open(ssot_file, 'w', encoding='utf-8') as f:
            yaml.dump(data, f, allow_unicode=True, sort_keys=False)
        
        # 透過 WebSocket 通知前端
        socketio.emit('ssot_updated', {'timestamp': datetime.now().isoformat()})
        
        return jsonify({
            'success': True,
            'message': 'SSOT 資料已更新'
        })
    except Exception as e:
        logger.error(f"更新 SSOT 失敗: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/ssot/flatten', methods=['GET'])
def get_ssot_flatten():
    """取得扁平化的 SSOT 資料 (用於欄位對應)"""
    try:
        ssot_file = SSOT_DIR / 'master.yaml'
        with open(ssot_file, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
        
        def flatten_dict(d, parent_key='', sep='.'):
            items = []
            for k, v in d.items():
                new_key = f"{parent_key}{sep}{k}" if parent_key else k
                if isinstance(v, dict):
                    items.extend(flatten_dict(v, new_key, sep=sep).items())
                elif isinstance(v, list):
                    # 列表項目不展開
                    items.append((new_key, v))
                else:
                    items.append((new_key, v))
            return dict(items)
        
        flattened = flatten_dict(data)
        
        return jsonify({
            'success': True,
            'data': flattened
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================================================
# API: 欄位對應管理
# ============================================================================

@app.route('/api/mapping', methods=['GET'])
def get_mapping():
    """讀取欄位對應設定"""
    try:
        mapping_file = MAPPING_DIR / 'customer_mapping.yaml'
        if not mapping_file.exists():
            return jsonify({'error': '對應表檔案不存在'}), 404
        
        with open(mapping_file, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
        
        return jsonify({
            'success': True,
            'data': data
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/mapping', methods=['POST'])
def update_mapping():
    """更新欄位對應設定"""
    try:
        data = request.get_json()
        mapping_file = MAPPING_DIR / 'customer_mapping.yaml'
        
        # 備份
        if mapping_file.exists():
            backup_file = MAPPING_DIR / f'customer_mapping.yaml.backup.{int(datetime.now().timestamp())}'
            import shutil
            shutil.copy(mapping_file, backup_file)
        
        # 更新時間戳
        data['last_updated'] = datetime.now().strftime('%Y-%m-%d')
        
        with open(mapping_file, 'w', encoding='utf-8') as f:
            yaml.dump(data, f, allow_unicode=True, sort_keys=False)
        
        socketio.emit('mapping_updated', {'timestamp': datetime.now().isoformat()})
        
        return jsonify({
            'success': True,
            'message': '對應表已更新'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================================================
# API: 模板管理
# ============================================================================

@app.route('/api/templates', methods=['GET'])
def list_templates():
    """列出所有模板檔案"""
    try:
        templates = []
        for file in TEMPLATES_DIR.iterdir():
            if file.is_file() and file.suffix in ['.docx', '.xlsx']:
                templates.append({
                    'name': file.name,
                    'type': 'Word' if file.suffix == '.docx' else 'Excel',
                    'size': file.stat().st_size,
                    'modified': datetime.fromtimestamp(file.stat().st_mtime).isoformat()
                })
        
        return jsonify({
            'success': True,
            'data': templates
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/templates/upload', methods=['POST'])
def upload_template():
    """上傳新模板"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': '未提供檔案'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '未選擇檔案'}), 400
        
        # 驗證檔案類型
        if not file.filename.endswith(('.docx', '.xlsx')):
            return jsonify({'error': '不支援的檔案類型'}), 400
        
        # 儲存檔案
        file_path = TEMPLATES_DIR / file.filename
        file.save(file_path)
        
        return jsonify({
            'success': True,
            'message': f'檔案 {file.filename} 上傳成功',
            'file': {
                'name': file.filename,
                'size': file_path.stat().st_size
            }
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================================================
# API: Token 掃描
# ============================================================================

@app.route('/api/templates/<path:filename>/scan', methods=['GET'])
def scan_template_tokens(filename):
    """掃描模板中的 Token，格式為 {path.to.value}
    回傳 tokens 列表與計數
    """
    try:
        file_path = TEMPLATES_DIR / filename
        if not file_path.exists():
            return jsonify({'error': '模板不存在'}), 404

        ext = file_path.suffix.lower()
        if ext == '.docx':
            try:
                tokens = sorted(_scan_tokens_docx(file_path))
            except ImportError:
                return jsonify({'error': '缺少套件 python-docx'}), 500
        elif ext == '.xlsx':
            try:
                tokens = sorted(_scan_tokens_xlsx(file_path))
            except ImportError:
                return jsonify({'error': '缺少套件 openpyxl'}), 500
        else:
            return jsonify({'error': '不支援的檔案類型'}), 400

        return jsonify({
            'success': True,
            'file': filename,
            'count': len(tokens),
            'tokens': tokens
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================================================
# API: 文件產生
# ============================================================================

@app.route('/api/generate', methods=['POST'])
def generate_documents():
    """產生文件：Token 優先，必要時回退到舊版腳本"""
    try:
        config = request.get_json() or {}
        engine = config.get('engine', 'auto')
        templates = config.get('templates', [])

        if not isinstance(templates, list) or not templates:
            return jsonify({'error': '請提供欲產生的模板清單'}), 400

        os.environ['SPEC_SYNC_ENGINE'] = engine

        socketio.emit('generate_start', {
            'timestamp': datetime.now().isoformat(),
            'templates': templates
        })

        results = []
        processed_success = set()
        token_mode_available = True
        token_success_count = 0

        # 嘗試 Token 模式
        try:
            ssot = _load_ssot()
        except Exception as e:
            # 無法載入 SSOT 時，無法執行 Token 替換
            token_mode_available = False
            ssot = None

        if token_mode_available:
            for template in templates:
                t_path = TEMPLATES_DIR / template
                if not t_path.exists():
                    results.append({
                        'template': template,
                        'status': 'error',
                        'error': '模板不存在'
                    })
                    continue

                ext = t_path.suffix.lower()
                try:
                    if ext == '.docx':
                        tokens = _scan_tokens_docx(t_path)
                    elif ext == '.xlsx':
                        tokens = _scan_tokens_xlsx(t_path)
                    else:
                        results.append({
                            'template': template,
                            'status': 'error',
                            'error': '不支援的檔案類型'
                        })
                        continue

                    if tokens:
                        out_name = f"filled_{template}"
                        out_path = OUTPUT_DIR / out_name
                        if ext == '.docx':
                            info = _replace_tokens_docx(t_path, out_path, ssot)
                        else:
                            info = _replace_tokens_xlsx(t_path, out_path, ssot)

                        token_success_count += 1
                        processed_success.add(template)
                        payload = {
                            'template': template,
                            'status': 'success',
                            'output': out_name,
                            'tokens_found': len(tokens),
                            'missing': info.get('missing', []),
                            'replaced_count': len(info.get('replaced', {}))
                        }
                        results.append(payload)
                        socketio.emit('generate_progress', payload)
                    else:
                        payload = {
                            'template': template,
                            'status': 'skipped',
                            'reason': '未找到 Token'
                        }
                        results.append(payload)
                        socketio.emit('generate_progress', payload)

                except ImportError as ie:
                    token_mode_available = False
                    break
                except Exception as e:
                    results.append({
                        'template': template,
                        'status': 'error',
                        'error': str(e)
                    })

        # 若 Token 模式不可用或沒有任何成功，回退舊版腳本
        if (not token_mode_available) or token_success_count == 0:
            script_path = project_root / 'scripts' / 'generate_docs.py'
            try:
                result = subprocess.run(
                    [sys.executable, str(script_path)],
                    cwd=str(project_root),
                    capture_output=True,
                    text=True,
                    timeout=300
                )

                if result.returncode == 0:
                    for template in templates:
                        if template in processed_success:
                            continue  # 已由 Token 模式產出
                        output_file = f"filled_{template}"
                        if (OUTPUT_DIR / output_file).exists():
                            payload = {
                                'template': template,
                                'status': 'success',
                                'output': output_file
                            }
                            results.append(payload)
                            socketio.emit('generate_progress', payload)
                        else:
                            results.append({
                                'template': template,
                                'status': 'error',
                                'error': 'Output file not found'
                            })
                else:
                    error_msg = result.stderr or result.stdout
                    for template in templates:
                        if template in processed_success:
                            continue
                        payload = {
                            'template': template,
                            'status': 'error',
                            'error': error_msg
                        }
                        results.append(payload)
                        socketio.emit('generate_progress', payload)

            except subprocess.TimeoutExpired:
                error_msg = 'Generation timeout (5 minutes)'
                for template in templates:
                    if template in processed_success:
                        continue
                    results.append({
                        'template': template,
                        'status': 'error',
                        'error': error_msg
                    })

        socketio.emit('generate_complete', {
            'timestamp': datetime.now().isoformat(),
            'results': results
        })

        return jsonify({
            'success': True,
            'results': results
        })

    except Exception as e:
        logger.error(f"Generate failed: {str(e)}")
        socketio.emit('generate_error', {
            'timestamp': datetime.now().isoformat(),
            'error': str(e)
        })
        return jsonify({'error': str(e)}), 500


@app.route('/api/validate', methods=['POST'])
def validate_documents():
    """Validate document consistency"""
    try:
        config = request.get_json()
        engine = config.get('engine', 'auto')
        
        os.environ['SPEC_SYNC_ENGINE'] = engine
        
        # Call existing validation script
        script_path = project_root / 'scripts' / 'validate_consistency.py'
        result = subprocess.run(
            [sys.executable, str(script_path)],
            cwd=str(project_root),
            capture_output=True,
            text=True,
            timeout=60
        )
        
        if result.returncode == 0:
            return jsonify({
                'success': True,
                'output': result.stdout,
                'message': 'Validation completed'
            })
        else:
            return jsonify({
                'success': False,
                'error': result.stderr or result.stdout
            }), 400
            
    except subprocess.TimeoutExpired:
        return jsonify({'error': 'Validation timeout'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================================================
# API: 檔案下載
# ============================================================================

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """下載產生的文件"""
    try:
        file_path = OUTPUT_DIR / filename
        if not file_path.exists():
            return jsonify({'error': '檔案不存在'}), 404
        
        return send_file(
            file_path,
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================================================
# API: 歷史記錄
# ============================================================================

HISTORY_FILE = OUTPUT_DIR / 'generation_history.json'

def save_history_record(record):
    """儲存歷史記錄"""
    history = []
    if HISTORY_FILE.exists():
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            history = json.load(f)
    
    history.insert(0, record)  # 最新的在前面
    history = history[:100]  # 保留最近 100 筆
    
    with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)


@app.route('/api/history', methods=['GET'])
def get_history():
    """取得產生歷史記錄"""
    try:
        if not HISTORY_FILE.exists():
            return jsonify({
                'success': True,
                'data': []
            })
        
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            history = json.load(f)
        
        return jsonify({
            'success': True,
            'data': history
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================================================
# API: 系統狀態
# ============================================================================

@app.route('/api/status', methods=['GET'])
def get_status():
    """取得系統狀態"""
    try:
        status = {
            'ssot_exists': (SSOT_DIR / 'master.yaml').exists(),
            'mapping_exists': (MAPPING_DIR / 'customer_mapping.yaml').exists(),
            'templates_count': len(list(TEMPLATES_DIR.glob('*.docx'))) + len(list(TEMPLATES_DIR.glob('*.xlsx'))),
            'output_count': len(list(OUTPUT_DIR.glob('*.docx'))) + len(list(OUTPUT_DIR.glob('*.xlsx'))),
            'python_version': sys.version,
            'server_time': datetime.now().isoformat()
        }
        
        return jsonify({
            'success': True,
            'data': status
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================================================
# WebSocket 事件
# ============================================================================

@socketio.on('connect')
def handle_connect():
    logger.info('客戶端已連接')
    emit('connected', {'message': '已連接到伺服器'})


@socketio.on('disconnect')
def handle_disconnect():
    logger.info('客戶端已斷開')


# ============================================================================
# 錯誤處理
# ============================================================================

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': '找不到資源'}), 404


@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': '伺服器內部錯誤'}), 500


# ============================================================================
# 主程式
# ============================================================================

if __name__ == '__main__':
    logger.info('啟動 Spec-Sync SSOT API 伺服器...')
    logger.info(f'專案根目錄: {project_root}')
    logger.info(f'SSOT 目錄: {SSOT_DIR}')
    logger.info(f'Templates 目錄: {TEMPLATES_DIR}')
    
    # 開發模式: http://localhost:5000
    socketio.run(app, host='0.0.0.0', port=5000, debug=True)
