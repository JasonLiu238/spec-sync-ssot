#!/usr/bin/env python3
"""
Spec Sync SSOT - 文件一致性驗證工具
檢查輸出文件與 SSOT 是否一致

支援純 Python 與 Office COM 兩種模式，與 generate_docs.py 相同。
"""

import os
import sys
import yaml
import json
import logging
from pathlib import Path
from typing import Dict, Any, List, Tuple

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def _import_python_doc_libs():
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        Document = None  # type: ignore
        logger.debug(f"python-docx 載入失敗: {e}")
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception as e:
        load_workbook = None  # type: ignore
        logger.debug(f"openpyxl 載入失敗: {e}")
    return Document, load_workbook

def _import_office_com():
    try:
        import win32com.client  # type: ignore
        return win32com.client
    except Exception as e:
        logger.debug(f"win32com 載入失敗: {e}")
        return None

class ConsistencyValidator:
    """文件一致性驗證器"""
    
    def __init__(self, base_path: str = "."):
        self.base_path = Path(base_path)
        self.ssot_path = self.base_path / "ssot"
        self.mapping_path = self.base_path / "mapping"
        self.output_path = self.base_path / "output"
        
    def load_ssot(self, ssot_file: str = "master.yaml") -> Dict[str, Any]:
        """載入 SSOT 檔案"""
        ssot_file_path = self.ssot_path / ssot_file
        
        with open(ssot_file_path, 'r', encoding='utf-8') as f:
            if ssot_file_path.suffix.lower() == '.yaml':
                return yaml.safe_load(f)
            elif ssot_file_path.suffix.lower() == '.json':
                return json.load(f)
                
    def load_mapping(self, mapping_file: str = "customer_mapping.yaml") -> Dict[str, Any]:
        """載入對應表"""
        mapping_file_path = self.mapping_path / mapping_file
        
        with open(mapping_file_path, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)
    
    def get_nested_value(self, data: Dict[str, Any], key_path: str) -> Any:
        """從巢狀字典中取得值"""
        keys = key_path.split('.')
        value = data
        
        try:
            for key in keys:
                value = value[key]
            return value
        except (KeyError, TypeError):
            return None
    
    def validate_word_document(self, doc_path: Path, mapping: Dict[str, str], 
                              ssot_data: Dict[str, Any]) -> List[str]:
        """驗證 Word 文件一致性（自動選擇引擎）。"""
        errors: List[str] = []
        if not doc_path.exists():
            return [f"Word 文件不存在: {doc_path}"]

        engine_pref = os.getenv("SPEC_SYNC_ENGINE", "auto").lower()
        Document, _ = _import_python_doc_libs()
        win32com = _import_office_com() if engine_pref in ("auto", "office") else None

        doc_text = None

        # 試 Python 解析
        if engine_pref in ("auto", "pure") and Document is not None:
            try:
                doc = Document(str(doc_path))
                text = []
                for p in doc.paragraphs:
                    text.append(p.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text.append(cell.text)
                doc_text = "\n".join(text)
            except Exception as e:
                logger.debug(f"python-docx 讀取失敗：{e}")

        # COM 讀取全文
        if doc_text is None and engine_pref in ("auto", "office") and win32com is not None:
            try:
                word = win32com.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(doc_path))
                doc_text = doc.Content.Text
                doc.Close(False)
                word.Quit()
            except Exception as e:
                logger.debug(f"Office Word 自動化讀取失敗：{e}")
                try:
                    word.Quit()
                except Exception:
                    pass

        if doc_text is None:
            return ["無法讀取 Word 文件內容（請確認權限或安裝必要套件）"]

        # 檢查每個對應欄位
        for ssot_field, _bookmark in mapping.items():
            expected_value = self.get_nested_value(ssot_data, ssot_field)
            if expected_value is None:
                continue
            if str(expected_value) not in doc_text:
                errors.append(
                    f"Word文件中找不到 {ssot_field} 的值: {expected_value}"
                )
        return errors
    
    def validate_excel_document(self, excel_path: Path, sheet_name: str,
                               mapping: Dict[str, str], ssot_data: Dict[str, Any]) -> List[str]:
        """驗證 Excel 文件一致性（自動選擇引擎）。"""
        errors: List[str] = []
        if not excel_path.exists():
            return [f"Excel 文件不存在: {excel_path}"]

        engine_pref = os.getenv("SPEC_SYNC_ENGINE", "auto").lower()
        _, load_workbook = _import_python_doc_libs()
        win32com = _import_office_com() if engine_pref in ("auto", "office") else None

        used_openpyxl = False
        if engine_pref in ("auto", "pure") and load_workbook is not None:
            try:
                wb = load_workbook(str(excel_path))
                if sheet_name not in wb.sheetnames:
                    return [f"工作表不存在: {sheet_name}"]
                ws = wb[sheet_name]
                for ssot_field, excel_cell in mapping.items():
                    expected_value = self.get_nested_value(ssot_data, ssot_field)
                    if expected_value is None:
                        continue
                    actual_value = ws[excel_cell].value
                    if str(actual_value) != str(expected_value):
                        errors.append(
                            f"Excel {excel_cell} 儲存格不一致: 期望 '{expected_value}', 實際 '{actual_value}'"
                        )
                used_openpyxl = True
            except Exception as e:
                logger.debug(f"openpyxl 讀取失敗：{e}")

        if not used_openpyxl and engine_pref in ("auto", "office") and win32com is not None:
            try:
                excel = win32com.Dispatch("Excel.Application")
                excel.Visible = False
                wb = excel.Workbooks.Open(str(excel_path))
                ws = wb.Worksheets(sheet_name)
                for ssot_field, excel_cell in mapping.items():
                    expected_value = self.get_nested_value(ssot_data, ssot_field)
                    if expected_value is None:
                        continue
                    actual_value = ws.Range(excel_cell).Value
                    if str(actual_value) != str(expected_value):
                        errors.append(
                            f"Excel {excel_cell} 儲存格不一致: 期望 '{expected_value}', 實際 '{actual_value}'"
                        )
                wb.Close(SaveChanges=False)
                excel.Quit()
            except Exception as e:
                errors.append(f"Office Excel 自動化讀取失敗：{e}")
                try:
                    excel.Quit()
                except Exception:
                    pass

        return errors
    
    def validate_all_documents(self) -> Tuple[bool, List[str]]:
        """驗證所有文件一致性"""
        all_errors = []
        
        try:
            # 載入 SSOT 和對應表
            ssot_data = self.load_ssot()
            mapping_config = self.load_mapping()
            
            logger.info("開始驗證文件一致性...")
            
            # 驗證 Word 文件
            if 'word_mappings' in mapping_config:
                for template_name, config in mapping_config['word_mappings'].items():
                    # 查找最新的輸出文件
                    pattern = f"{template_name}_*.docx"
                    matching_files = list(self.output_path.glob(pattern))
                    
                    if not matching_files:
                        all_errors.append(f"找不到 {template_name} 的輸出文件")
                        continue
                    
                    # 取最新的檔案
                    latest_file = max(matching_files, key=lambda x: x.stat().st_mtime)
                    
                    errors = self.validate_word_document(
                        latest_file, 
                        config['mappings'], 
                        ssot_data
                    )
                    all_errors.extend(errors)
            
            # 驗證 Excel 文件
            if 'excel_mappings' in mapping_config:
                for template_name, config in mapping_config['excel_mappings'].items():
                    # 查找最新的輸出文件
                    pattern = f"{template_name}_*.xlsx"
                    matching_files = list(self.output_path.glob(pattern))
                    
                    if not matching_files:
                        all_errors.append(f"找不到 {template_name} 的輸出文件")
                        continue
                    
                    # 取最新的檔案
                    latest_file = max(matching_files, key=lambda x: x.stat().st_mtime)
                    
                    errors = self.validate_excel_document(
                        latest_file,
                        config.get('sheet_name', 'Sheet1'),
                        config['mappings'], 
                        ssot_data
                    )
                    all_errors.extend(errors)
            
            is_valid = len(all_errors) == 0
            
            if is_valid:
                logger.info("✅ 所有文件驗證通過")
            else:
                logger.error(f"❌ 發現 {len(all_errors)} 個一致性錯誤")
                
            return is_valid, all_errors
            
        except Exception as e:
            error_msg = f"驗證過程中發生錯誤: {e}"
            logger.error(error_msg)
            return False, [error_msg]

def main():
    """主程式入口"""
    validator = ConsistencyValidator()
    
    is_valid, errors = validator.validate_all_documents()
    
    if errors:
        print("\n❌ 發現以下一致性問題:")
        for i, error in enumerate(errors, 1):
            print(f"{i}. {error}")
    
    if is_valid:
        print("\n✅ 所有文件與 SSOT 一致！")
        sys.exit(0)
    else:
        print(f"\n❌ 驗證失敗，發現 {len(errors)} 個問題")
        sys.exit(1)

if __name__ == "__main__":
    main()