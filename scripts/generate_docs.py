#!/usr/bin/env python3
"""
Spec Sync SSOT - 文件自動產生引擎
從 SSOT 資料自動填寫客戶 Word 和 Excel 模板

支援兩種模式：
1) 純 Python（python-docx / openpyxl）
2) Office COM 自動化（win32com）→ 可處理受敏感性標籤/IRM 保護的文件（需權限）

以環境變數 SPEC_SYNC_ENGINE 控制：auto | pure | office（預設 auto）
"""

import os
import sys
import yaml
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any

# 設定日誌
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def _import_python_doc_libs():
    """延遲載入 python-docx 與 openpyxl。"""
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        Document = None  # type: ignore
        logger.debug(f"python-docx 載入失敗: {e}")
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception as e:
        load_workbook = None  # type: ignore
        logger.debug(f"openpyxl 載入失敗: {e}")
    return Document, load_workbook

def _import_office_com():
    """延遲載入 win32com。"""
    try:
        import win32com.client  # type: ignore
        return win32com.client
    except Exception as e:
        logger.debug(f"win32com 載入失敗: {e}")
        return None

class SpecSyncEngine:
    """規格同步引擎"""
    
    def __init__(self, base_path: str = "."):
        self.base_path = Path(base_path)
        self.ssot_path = self.base_path / "ssot"
        self.mapping_path = self.base_path / "mapping" 
        self.template_path = self.base_path / "templates"
        self.output_path = self.base_path / "output"
        
        # 確保輸出目錄存在
        self.output_path.mkdir(exist_ok=True)
        
    def load_ssot(self, ssot_file: str = "master.yaml") -> Dict[str, Any]:
        """載入 SSOT 主檔案"""
        ssot_file_path = self.ssot_path / ssot_file
        
        if not ssot_file_path.exists():
            raise FileNotFoundError(f"SSOT 檔案不存在: {ssot_file_path}")
            
        with open(ssot_file_path, 'r', encoding='utf-8') as f:
            if ssot_file_path.suffix.lower() == '.yaml':
                return yaml.safe_load(f)
            elif ssot_file_path.suffix.lower() == '.json':
                return json.load(f)
        
        raise ValueError(f"不支援的 SSOT 檔案格式: {ssot_file_path.suffix}")
    
    def load_mapping(self, mapping_file: str = "customer_mapping.yaml") -> Dict[str, Any]:
        """載入客戶欄位對應表"""
        mapping_file_path = self.mapping_path / mapping_file
        
        if not mapping_file_path.exists():
            raise FileNotFoundError(f"對應表檔案不存在: {mapping_file_path}")
            
        with open(mapping_file_path, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)
    
    def get_nested_value(self, data: Dict[str, Any], key_path: str) -> Any:
        """從巢狀字典中取得值 (例如: product.name -> data['product']['name'])"""
        keys = key_path.split('.')
        value = data
        
        try:
            for key in keys:
                value = value[key]
            return value
        except (KeyError, TypeError):
            logger.warning(f"找不到欄位: {key_path}")
            return None
    
    def fill_word_template(self, template_file: str, mapping: Dict[str, str], 
                          ssot_data: Dict[str, Any], output_file: str):
        """填寫 Word 模板（自動選擇引擎）。"""
        engine_pref = os.getenv("SPEC_SYNC_ENGINE", "auto").lower()
        template_path = self.template_path / template_file
        output_path = self.output_path / output_file

        if not template_path.exists():
            logger.error(f"Word 模板不存在: {template_path}")
            return False

        Document, _ = _import_python_doc_libs()
        win32com = _import_office_com() if engine_pref in ("auto", "office") else None

        def _fill_with_python_docx() -> bool:
            if Document is None:
                return False
            try:
                doc = Document(str(template_path))
                # 替換書籤/欄位（以 Token {Bookmark} 為主）
                for ssot_field, word_bookmark in mapping.items():
                    value = self.get_nested_value(ssot_data, ssot_field)
                    if value is None:
                        continue
                    token = f"{{{word_bookmark}}}"
                    for paragraph in doc.paragraphs:
                        if token in paragraph.text:
                            paragraph.text = paragraph.text.replace(token, str(value))
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if token in cell.text:
                                    cell.text = cell.text.replace(token, str(value))
                doc.save(str(output_path))
                logger.info(f"Word 文件已產生: {output_path}")
                return True
            except Exception as e:
                logger.warning(f"python-docx 處理失敗，將嘗試 Office 模式：{e}")
                return False

        def _fill_with_office_com() -> bool:
            if win32com is None:
                return False
            try:
                word = win32com.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(template_path))
                # 嘗試書籤填入；若無則使用尋找取代
                for ssot_field, name in mapping.items():
                    value = self.get_nested_value(ssot_data, ssot_field)
                    if value is None:
                        continue
                    try:
                        if doc.Bookmarks.Exists(name):
                            doc.Bookmarks(name).Range.Text = str(value)
                            continue
                    except Exception:
                        pass
                    # Find/Replace token {Name} 於整份文件
                    rng = doc.Content
                    find = rng.Find
                    find.ClearFormatting()
                    find.Text = f"{{{name}}}"
                    find.Replacement.ClearFormatting()
                    find.Replacement.Text = str(value)
                    find.Execute(Replace=2)  # wdReplaceAll
                # 另存新檔為 .docx
                wdFormatXMLDocument = 12
                doc.SaveAs(str(output_path), FileFormat=wdFormatXMLDocument)
                doc.Close(False)
                word.Quit()
                logger.info(f"Word 文件已產生（Office 模式）: {output_path}")
                return True
            except Exception as e:
                logger.error(f"Office Word 自動化失敗：{e}")
                try:
                    # 嘗試確保應用程式關閉
                    word.Quit()
                except Exception:
                    pass
                return False

        # 根據偏好選擇
        if engine_pref == "pure":
            return _fill_with_python_docx()
        elif engine_pref == "office":
            return _fill_with_office_com()
        else:
            return _fill_with_python_docx() or _fill_with_office_com()
    
    def fill_excel_template(self, template_file: str, sheet_name: str,
                           mapping: Dict[str, str], ssot_data: Dict[str, Any], 
                           output_file: str):
        """填寫 Excel 模板（自動選擇引擎）。"""
        engine_pref = os.getenv("SPEC_SYNC_ENGINE", "auto").lower()
        template_path = self.template_path / template_file
        output_path = self.output_path / output_file

        if not template_path.exists():
            logger.error(f"Excel 模板不存在: {template_path}")
            return False

        _, load_workbook = _import_python_doc_libs()
        win32com = _import_office_com() if engine_pref in ("auto", "office") else None

        def _fill_with_openpyxl() -> bool:
            if load_workbook is None:
                return False
            try:
                wb = load_workbook(str(template_path))
                if sheet_name not in wb.sheetnames:
                    logger.error(f"工作表不存在: {sheet_name}")
                    return False
                ws = wb[sheet_name]
                for ssot_field, excel_cell in mapping.items():
                    value = self.get_nested_value(ssot_data, ssot_field)
                    if value is not None:
                        ws[excel_cell] = value
                wb.save(str(output_path))
                logger.info(f"Excel 文件已產生: {output_path}")
                return True
            except Exception as e:
                logger.warning(f"openpyxl 處理失敗，將嘗試 Office 模式：{e}")
                return False

        def _fill_with_office_com() -> bool:
            if win32com is None:
                return False
            try:
                excel = win32com.Dispatch("Excel.Application")
                excel.Visible = False
                wb = excel.Workbooks.Open(str(template_path))
                ws = wb.Worksheets(sheet_name)
                for ssot_field, excel_cell in mapping.items():
                    value = self.get_nested_value(ssot_data, ssot_field)
                    if value is not None:
                        ws.Range(excel_cell).Value = value
                # 另存新檔為 .xlsx
                xlOpenXMLWorkbook = 51
                wb.SaveAs(str(output_path), FileFormat=xlOpenXMLWorkbook)
                wb.Close(SaveChanges=False)
                excel.Quit()
                logger.info(f"Excel 文件已產生（Office 模式）: {output_path}")
                return True
            except Exception as e:
                logger.error(f"Office Excel 自動化失敗：{e}")
                try:
                    excel.Quit()
                except Exception:
                    pass
                return False

        if engine_pref == "pure":
            return _fill_with_openpyxl()
        elif engine_pref == "office":
            return _fill_with_office_com()
        else:
            return _fill_with_openpyxl() or _fill_with_office_com()
    
    def generate_all_documents(self):
        """產生所有文件"""
        try:
            # 載入 SSOT 和對應表
            ssot_data = self.load_ssot()
            mapping_config = self.load_mapping()
            
            logger.info("開始產生客戶文件...")
            
            # 處理 Word 文件
            if 'word_mappings' in mapping_config:
                for template_name, config in mapping_config['word_mappings'].items():
                    template_file = config['file_path'].replace('templates/', '')
                    output_file = f"{template_name}_{datetime.now().strftime('%Y%m%d')}.docx"
                    
                    self.fill_word_template(
                        template_file, 
                        config['mappings'], 
                        ssot_data, 
                        output_file
                    )
            
            # 處理 Excel 文件
            if 'excel_mappings' in mapping_config:
                for template_name, config in mapping_config['excel_mappings'].items():
                    template_file = config['file_path'].replace('templates/', '')
                    sheet_name = config.get('sheet_name', 'Sheet1')
                    output_file = f"{template_name}_{datetime.now().strftime('%Y%m%d')}.xlsx"
                    
                    self.fill_excel_template(
                        template_file,
                        sheet_name,
                        config['mappings'], 
                        ssot_data, 
                        output_file
                    )
            
            logger.info("所有文件產生完成！")
            return True
            
        except Exception as e:
            logger.error(f"產生文件時發生錯誤: {e}")
            return False

def main():
    """主程式入口"""
    engine = SpecSyncEngine()
    
    if engine.generate_all_documents():
        print("✅ 文件產生成功！請檢查 output/ 資料夾")
        sys.exit(0)
    else:
        print("❌ 文件產生失敗，請檢查日誌")
        sys.exit(1)

if __name__ == "__main__":
    main()