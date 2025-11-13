# ==============================================================================
# 工具 1: Word 文件自動標記輔助工具
# 檔案名稱: auto_bookmark_helper.py
# 用途: 掃描 Word 文件,找出所有可能需要標記的欄位,並建議書籤名稱
# ==============================================================================

import os
import re
from pathlib import Path
import yaml

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx 未安裝,嘗試使用 COM 模式...")

try:
    import win32com.client
    COM_AVAILABLE = True
except ImportError:
    COM_AVAILABLE = False


def extract_potential_fields_from_docx(file_path):
    """
    從 Word 文件中提取可能需要標記的欄位
    使用啟發式規則識別:
    1. 冒號後的內容 (例如: "產品名稱: _____")
    2. 表格中的空白儲存格
    3. 特定格式的文字 (例如: [待填入])
    """
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document(file_path)
    potential_fields = []
    
    # 規則 1: 尋找 "欄位名稱: _____" 或 "欄位名稱: [空白]" 格式
    pattern_colon = re.compile(r'([^\n:：]+)[：:]\s*(_+|\[.*?\]|【.*?】|＿+|\s{3,}|$)')
    
    # 規則 2: 尋找常見欄位關鍵字
    field_keywords = [
        '名稱', '版本', '型號', '規格', '描述', '說明',
        'CPU', '記憶體', '硬碟', '儲存', '作業系統', 
        '日期', '時間', '預算', '金額', '數量',
        '負責人', '聯絡', '電話', '地址', 'Email'
    ]
    
    paragraph_index = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        paragraph_index += 1
        
        # 檢查冒號格式
        matches = pattern_colon.findall(text)
        for field_name, placeholder in matches:
            field_name = field_name.strip()
            if len(field_name) < 30:  # 避免抓到太長的句子
                potential_fields.append({
                    'type': 'paragraph',
                    'location': f'段落 {paragraph_index}',
                    'field_name': field_name,
                    'context': text[:100],
                    'suggested_bookmark': generate_bookmark_name(field_name),
                    'confidence': 'high' if placeholder else 'medium'
                })
        
        # 檢查是否包含關鍵字
        for keyword in field_keywords:
            if keyword in text and len(text) < 50:
                if not any(pf['field_name'] == text for pf in potential_fields):
                    potential_fields.append({
                        'type': 'keyword',
                        'location': f'段落 {paragraph_index}',
                        'field_name': text,
                        'context': text,
                        'suggested_bookmark': generate_bookmark_name(text),
                        'confidence': 'medium'
                    })
    
    # 檢查表格
    table_index = 0
    for table in doc.tables:
        table_index += 1
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # 表格欄位名稱通常在第一列或第一欄
                if row_idx == 0 or col_idx == 0:
                    if cell_text and any(kw in cell_text for kw in field_keywords):
                        # 找對應的值儲存格
                        value_cell = None
                        if col_idx == 0 and len(row.cells) > 1:
                            value_cell = row.cells[1].text.strip()
                        elif row_idx == 0 and table_index < len(table.rows):
                            value_cell = table.rows[row_idx + 1].cells[col_idx].text.strip()
                        
                        if not value_cell or len(value_cell) < 3:  # 空白或很短 = 可能需要填入
                            potential_fields.append({
                                'type': 'table',
                                'location': f'表格 {table_index}, 列 {row_idx + 1}, 欄 {col_idx + 1}',
                                'field_name': cell_text,
                                'context': cell_text,
                                'suggested_bookmark': generate_bookmark_name(cell_text),
                                'confidence': 'high'
                            })
    
    return potential_fields


def generate_bookmark_name(field_name):
    """
    從中文欄位名稱生成英文書籤名稱
    """
    # 預定義常見對應
    mapping = {
        '產品名稱': 'ProductName',
        '產品型號': 'ProductModel',
        '版本': 'Version',
        '版本號': 'VersionNumber',
        '描述': 'Description',
        '說明': 'Description',
        'CPU': 'CPU',
        '處理器': 'CPU',
        '記憶體': 'Memory',
        'RAM': 'Memory',
        '硬碟': 'Storage',
        '儲存空間': 'Storage',
        '作業系統': 'OS',
        '開始日期': 'StartDate',
        '結束日期': 'EndDate',
        '預算': 'Budget',
        '金額': 'Amount',
        '負責人': 'Owner',
        '聯絡人': 'Contact',
        '電話': 'Phone',
        '地址': 'Address',
        '郵件': 'Email',
        'Email': 'Email',
    }
    
    # 先嘗試直接對應
    if field_name in mapping:
        return mapping[field_name]
    
    # 移除常見後綴
    clean_name = field_name.replace('名稱', '').replace('編號', 'ID').strip()
    if clean_name in mapping:
        return mapping[clean_name]
    
    # 轉換為拼音或保持英文
    # 這裡簡化處理,實際可整合 pypinyin 套件
    # 如果包含英文,保留英文
    english_only = re.sub(r'[^a-zA-Z0-9]', '', field_name)
    if english_only:
        return english_only[:30]  # 限制長度
    
    # 轉為拼音首字母 (簡化版,建議使用 pypinyin)
    return 'Field_' + ''.join(filter(str.isalnum, field_name))[:20]


def generate_mapping_suggestions(potential_fields, ssot_path='ssot/master.yaml'):
    """
    根據 SSOT 結構,建議欄位對應
    """
    if not os.path.exists(ssot_path):
        return potential_fields
    
    with open(ssot_path, 'r', encoding='utf-8') as f:
        ssot = yaml.safe_load(f)
    
    # 扁平化 SSOT 結構
    ssot_fields = flatten_dict(ssot)
    
    # 為每個潛在欄位找最佳匹配
    for field in potential_fields:
        field_lower = field['field_name'].lower()
        best_match = None
        best_score = 0
        
        for ssot_key, ssot_value in ssot_fields.items():
            # 簡單相似度計算
            score = calculate_similarity(field_lower, ssot_key.lower())
            if score > best_score:
                best_score = score
                best_match = ssot_key
        
        if best_score > 0.3:  # 相似度閾值
            field['suggested_ssot_path'] = best_match
            field['ssot_value'] = ssot_fields[best_match]
        else:
            field['suggested_ssot_path'] = None
    
    return potential_fields


def flatten_dict(d, parent_key='', sep='.'):
    """扁平化嵌套字典"""
    items = []
    for k, v in d.items():
        new_key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            items.extend(flatten_dict(v, new_key, sep=sep).items())
        elif isinstance(v, list):
            # 跳過列表
            continue
        else:
            items.append((new_key, v))
    return dict(items)


def calculate_similarity(text1, text2):
    """簡單的文字相似度計算"""
    # 檢查關鍵字匹配
    keywords_map = {
        'name': ['名稱', 'name'],
        'version': ['版本', 'version'],
        'cpu': ['cpu', '處理器', 'processor'],
        'memory': ['記憶體', 'memory', 'ram'],
        'storage': ['硬碟', '儲存', 'storage', 'disk'],
        'os': ['作業系統', 'os', 'operating'],
    }
    
    for key, keywords in keywords_map.items():
        if any(kw in text1 for kw in keywords) and any(kw in text2 for kw in keywords):
            return 0.8
    
    # 簡單字串包含
    if text1 in text2 or text2 in text1:
        return 0.5
    
    return 0.0


def export_to_yaml(potential_fields, template_name, output_path='mapping/auto_generated_mapping.yaml'):
    """
    將建議的欄位對應匯出為 YAML 格式
    """
    mapping = {
        'mapping_version': '1.0.0',
        'last_updated': '2025-11-13',
        'word_mappings': {
            template_name: {
                'file_path': f'templates/{template_name}.docx',
                'mappings': {}
            }
        }
    }
    
    for field in potential_fields:
        if field.get('suggested_ssot_path'):
            bookmark = field['suggested_bookmark']
            ssot_path = field['suggested_ssot_path']
            mapping['word_mappings'][template_name]['mappings'][ssot_path] = bookmark
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        yaml.dump(mapping, f, allow_unicode=True, sort_keys=False)
    
    return output_path


def generate_report(potential_fields, output_path='output/bookmark_suggestions.txt'):
    """
    產生人類可讀的報告
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("Word 文件自動標記建議報告\n")
        f.write("=" * 80 + "\n\n")
        
        # 按信心度分組
        high_conf = [f for f in potential_fields if f['confidence'] == 'high']
        medium_conf = [f for f in potential_fields if f['confidence'] == 'medium']
        
        f.write(f"📊 統計資訊:\n")
        f.write(f"  • 總共找到 {len(potential_fields)} 個潛在欄位\n")
        f.write(f"  • 高信心度: {len(high_conf)} 個\n")
        f.write(f"  • 中信心度: {len(medium_conf)} 個\n")
        f.write(f"  • 已建議 SSOT 對應: {len([f for f in potential_fields if f.get('suggested_ssot_path')])} 個\n\n")
        
        f.write("=" * 80 + "\n")
        f.write("高信心度欄位 (建議優先標記)\n")
        f.write("=" * 80 + "\n\n")
        
        for idx, field in enumerate(high_conf, 1):
            f.write(f"{idx}. {field['field_name']}\n")
            f.write(f"   位置: {field['location']}\n")
            f.write(f"   類型: {field['type']}\n")
            f.write(f"   建議書籤名稱: {field['suggested_bookmark']}\n")
            if field.get('suggested_ssot_path'):
                f.write(f"   建議 SSOT 路徑: {field['suggested_ssot_path']}\n")
                f.write(f"   目前 SSOT 值: {field.get('ssot_value', 'N/A')}\n")
            f.write(f"   上下文: {field['context']}\n")
            f.write("\n")
        
        if medium_conf:
            f.write("=" * 80 + "\n")
            f.write("中信心度欄位 (請手動確認)\n")
            f.write("=" * 80 + "\n\n")
            
            for idx, field in enumerate(medium_conf, 1):
                f.write(f"{idx}. {field['field_name']}\n")
                f.write(f"   位置: {field['location']}\n")
                f.write(f"   建議書籤名稱: {field['suggested_bookmark']}\n")
                if field.get('suggested_ssot_path'):
                    f.write(f"   建議 SSOT 路徑: {field['suggested_ssot_path']}\n")
                f.write("\n")
    
    return output_path


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='Word 文件自動標記輔助工具')
    parser.add_argument('file_path', help='Word 文件路徑')
    parser.add_argument('--template-name', default='auto_detected', help='模板名稱')
    parser.add_argument('--ssot', default='ssot/master.yaml', help='SSOT 檔案路徑')
    parser.add_argument('--output-report', default='output/bookmark_suggestions.txt', help='報告輸出路徑')
    parser.add_argument('--output-mapping', default='mapping/auto_generated_mapping.yaml', help='對應表輸出路徑')
    
    args = parser.parse_args()
    
    print("=" * 80)
    print("Word 文件自動標記輔助工具")
    print("=" * 80)
    print()
    
    if not DOCX_AVAILABLE and not COM_AVAILABLE:
        print("❌ 錯誤: 需要安裝 python-docx 或 pywin32")
        print("   pip install python-docx")
        return
    
    if not os.path.exists(args.file_path):
        print(f"❌ 錯誤: 找不到檔案 {args.file_path}")
        return
    
    print(f"📂 分析文件: {args.file_path}")
    print()
    
    try:
        # 提取潛在欄位
        print("🔍 掃描文件,尋找潛在欄位...")
        potential_fields = extract_potential_fields_from_docx(args.file_path)
        
        if not potential_fields:
            print("⚠️  未找到潛在欄位")
            return
        
        print(f"✅ 找到 {len(potential_fields)} 個潛在欄位")
        print()
        
        # 建議 SSOT 對應
        print("🔗 分析 SSOT 對應...")
        potential_fields = generate_mapping_suggestions(potential_fields, args.ssot)
        matched = len([f for f in potential_fields if f.get('suggested_ssot_path')])
        print(f"✅ 建議了 {matched} 個 SSOT 對應")
        print()
        
        # 產生報告
        print("📝 產生報告...")
        report_path = generate_report(potential_fields, args.output_report)
        print(f"✅ 報告已儲存: {report_path}")
        print()
        
        # 匯出對應表
        print("💾 匯出對應表...")
        mapping_path = export_to_yaml(potential_fields, args.template_name, args.output_mapping)
        print(f"✅ 對應表已儲存: {mapping_path}")
        print()
        
        print("=" * 80)
        print("✅ 完成!")
        print("=" * 80)
        print()
        print("📋 下一步:")
        print(f"  1. 查看報告: {report_path}")
        print(f"  2. 檢查對應表: {mapping_path}")
        print("  3. 在 Word 中手動建立書籤 (或使用批次建立工具)")
        print("  4. 執行文件產生測試")
        
    except Exception as e:
        print(f"❌ 錯誤: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    main()
