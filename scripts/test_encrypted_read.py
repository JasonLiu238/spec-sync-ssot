#!/usr/bin/env python3
"""測試加密檔案讀取能力"""
from docx import Document
import sys
from pathlib import Path

doc_path = Path("d:/AI/spec-sync-ssot/templates/customer_template_1.docx")

print("=" * 60)
print("測試 python-docx 讀取加密檔案")
print("=" * 60)
print(f"檔案: {doc_path}")
print(f"檔案存在: {doc_path.exists()}")
print(f"檔案大小: {doc_path.stat().st_size if doc_path.exists() else 0} bytes")
print()

try:
    doc = Document(str(doc_path))
    print("✅ 成功開啟文件！")
    print()
    print(f"📊 文件結構:")
    print(f"  - 段落數量: {len(doc.paragraphs)}")
    print(f"  - 表格數量: {len(doc.tables)}")
    print(f"  - 樣式數量: {len(doc.styles)}")
    print()
    
    print("📝 前 5 個段落內容:")
    for i, p in enumerate(doc.paragraphs[:5], 1):
        text = p.text.strip()
        if text:
            preview = text[:60] + "..." if len(text) > 60 else text
            print(f"  {i}. {preview}")
    print()
    
    if doc.tables:
        print("📋 第一個表格:")
        table = doc.tables[0]
        print(f"  - 行數: {len(table.rows)}")
        print(f"  - 列數: {len(table.columns) if table.rows else 0}")
    
    print()
    print("✅ python-docx 可以讀取此檔案（非加密或權限允許）")
    sys.exit(0)
    
except Exception as e:
    print(f"❌ 讀取失敗")
    print(f"錯誤: {e}")
    print(f"錯誤類型: {type(e).__name__}")
    print()
    print("📌 此檔案可能:")
    print("  1. 受 IRM/敏感性標籤保護")
    print("  2. 需要特定權限才能開啟")
    print("  3. 使用 COM 自動化或 VBA 巨集替代方案")
    sys.exit(1)
